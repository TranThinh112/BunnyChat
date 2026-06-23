from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
DOCS_DIR = ROOT / "Frontend" / "wwwroot" / "docs"
SHOT_DIR = DOCS_DIR / "report-shots"
OUTPUT = DOCS_DIR / "BaoCao_DoAn_BunnyChat.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
BORDER = "D9E2F3"
TEXT = "222222"
MUTED = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=False, color=TEXT, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_font(run)
    else:
        run = paragraph.add_run(text)
        set_font(run)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_font(run)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_font(run)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_font(run, size=9, color=MUTED)
    return paragraph


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_border(cell)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    r = p.add_run(body)
    set_font(r)
    doc.add_paragraph()


def add_kv_table(doc, rows, widths=(1.8, 4.7)):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, list(widths))
    for i, (key, value) in enumerate(rows):
        left, right = table.rows[i].cells
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, LIGHT_GRAY)
        left.text = ""
        right.text = ""
        r = left.paragraphs[0].add_run(key)
        set_font(r, bold=True, color=DARK_BLUE)
        r = right.paragraphs[0].add_run(value)
        set_font(r)
    doc.add_paragraph()
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = cell.paragraphs[0].add_run(header)
        set_font(r, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = cell.paragraphs[0].add_run(value)
            set_font(r, size=10)
    doc.add_paragraph()
    return table


def add_image(doc, filename, caption, width=6.1):
    path = SHOT_DIR / filename
    if not path.exists():
        add_callout(doc, "Anh minh hoa chua co", f"Khong tim thay file: {path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "BunnyChat - Bao cao do an"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Bao cao he thong realtime chat - toi da 20 trang")
    set_font(run, size=9, color=MUTED)
    return doc


def build_report():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BAO CAO DO AN\nBUNNYCHAT REALTIME CHAT APP")
    set_font(run, size=22, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("ASP.NET Core MVC + Web API + SignalR + MongoDB")
    set_font(run, size=13, bold=True, color=BLUE)

    doc.add_paragraph()
    add_kv_table(doc, [
        ("Ten do an", "BunnyChat - ung dung chat realtime chu de tho"),
        ("Cong nghe", "ASP.NET Core 8, Razor/MVC, Web API, SignalR, MongoDB, JavaScript, HTML/CSS"),
        ("Pham vi", "Dang nhap, bao ve route, chat realtime, danh sach hoi thoai, tim user, ket ban, thong bao loi moi, graph he thong"),
        ("Ngay lap bao cao", "23/06/2026"),
    ])

    add_callout(
        doc,
        "Tom tat",
        "Do an BunnyChat xay dung mot he thong chat realtime co giao dien chu de tho, co frontend ASP.NET MVC, backend API, MongoDB de luu du lieu va SignalR de day su kien realtime cho tin nhan va ket ban."
    )

    add_heading(doc, "Muc luc ngan", 1)
    add_numbered(doc, [
        "Gioi thieu va muc tieu do an",
        "Cac chuc nang da hoan thanh",
        "Kien truc tong quan va anh minh hoa",
        "Luong logic chat realtime va ket ban",
        "API, SignalR event va response chinh",
        "Cau truc file va huong phat trien",
    ])

    doc.add_page_break()

    add_heading(doc, "1. Gioi Thieu Do An", 1)
    add_paragraph(
        doc,
        "BunnyChat la ung dung chat realtime duoc thiet ke theo chu de tho, tap trung vao trai nghiem chat nhom, chat truc tiep, tim kiem nguoi dung va ket ban. Trong qua trinh hoan thien, du lieu mau tren UI da duoc loai bo de frontend render theo data server tra ve."
    )
    add_paragraph(
        doc,
        "He thong duoc chia thanh hai lop chinh: frontend phu trach giao dien, tuong tac nguoi dung va ket noi realtime; backend phu trach xac thuc, xu ly API, luu tru MongoDB va phat event SignalR."
    )
    add_heading(doc, "Muc tieu", 2)
    add_bullets(doc, [
        "Xay dung giao dien chat hien dai, co dark/light theme va animation chu de tho.",
        "Cho phep frontend lay du lieu hoi thoai, ban be, loi moi ket ban tu backend thay vi dung data co dinh.",
        "Bo sung SignalR de cap nhat tin nhan va trang thai ket ban theo thoi gian thuc.",
        "Giu kien truc cu, chi bo sung API/chuc nang con thieu va khong lam thay doi logic API cu.",
    ])

    add_heading(doc, "2. Cac Chuc Nang Da Hoan Thanh", 1)
    add_matrix_table(
        doc,
        ["Nhom chuc nang", "Noi dung da lam", "Trang thai"],
        [
            ("Giao dien chat", "Trang /Chat co sidebar, group chat, friends, message area, input gui tin nhan, theme sang/toi.", "Da hoan thanh UI"),
            ("Chu de tho", "Logo tho, toggle theme tho trang/tho den, animation tai tho khi hover card.", "Da hoan thanh"),
            ("Tim kiem user", "Search user theo keyword, hien card ten va username, bam card xem thong tin profile.", "Da ket noi API"),
            ("Ket ban", "Nut Ket ban nam trong card user search, co loi moi da gui/da nhan, accept/decline.", "Da bo sung UI + API"),
            ("Thong bao", "Chuong thong bao goc trai, badge dem loi moi ket ban moi.", "Da bo sung"),
            ("Realtime", "SignalR cho new-message, new-conversation, read-message va friend request events.", "Da bo sung"),
            ("Graph he thong", "File HTML graph mo ta kien truc, API, response, luong logic va file map.", "Da bo sung"),
        ],
        [1.45, 3.95, 1.1],
    )

    add_heading(doc, "3. Anh Minh Hoa Chuc Nang", 1)
    add_paragraph(doc, "Cac anh ben duoi duoc chup tu he thong dang chay local va file graph chi tiet cua do an.")
    add_image(doc, "06-auth-login.png", "Hinh 1. Man hinh dang nhap va dang ky cua BunnyChat.", 5.9)
    add_image(doc, "01-graph-overview.png", "Hinh 2. Tong quan kien truc BunnyChat.", 6.1)

    doc.add_page_break()

    add_heading(doc, "4. Kien Truc Tong Quan", 1)
    add_paragraph(
        doc,
        "Kien truc hien tai di theo luong: View /Chat render HTML Razor, ChatAction.js dieu phoi UI va goi API, Api_Fetch.js gan token, Controller xu ly request, ChatHub day event realtime, MongoDB luu du lieu."
    )
    add_kv_table(doc, [
        ("Frontend View", "Frontend/Views/Chat/Index.cshtml render layout chat, chuong thong bao, modal profile va modal ket ban."),
        ("Frontend JS", "Frontend/wwwroot/js/ChatAPI/ChatAction.js quan ly state, render UI, goi API va lang nghe SignalR."),
        ("Backend API", "ChatController, FriendController, UserController tra response cho UI va phat event realtime khi can."),
        ("Realtime", "ChatHub quan ly group user:{id} va conversation:{id}."),
        ("Database", "MongoDB luu users, friends, conversations va messages."),
    ])
    add_image(doc, "02-function-graph.png", "Hinh 3. Function graph: ham frontend goi API va cap nhat UI.", 6.1)

    add_heading(doc, "5. Luong Logic Hoat Dong", 1)
    add_heading(doc, "5.1. Luong mo trang chat", 2)
    add_numbered(doc, [
        "Nguoi dung truy cap /Chat.",
        "Protect_Chat.js kiem tra token. Neu chua dang nhap hoac token khong hop le thi redirect ve trang Auth.",
        "loadChatData() goi /api/users/me, /api/chat, /api/friends va /api/friends/requests.",
        "Frontend cap nhat state.currentUser, state.conversations, state.friends va state.friendRequests.",
        "renderGroups(), renderFriends() va renderFriendRequestBadge() render lai sidebar va chuong thong bao.",
        "connectRealtime() mo ket noi SignalR toi /chatHub.",
    ])

    add_heading(doc, "5.2. Luong gui tin nhan realtime", 2)
    add_numbered(doc, [
        "Nguoi dung bam vao conversation hoac friend card.",
        "selectConversation() goi GET /api/chat/{conversationId}/messages de lay lich su tin nhan.",
        "joinConversation() dua client vao SignalR group conversation:{id}.",
        "Khi gui tin nhan, form goi POST /api/chat/{id}/messages.",
        "ChatController.SendMessage() luu message, update last message va unread count.",
        "Server phat event new-message cho group conversation; client nhan event va appendRealtimeMessage() them bubble moi.",
    ])

    add_heading(doc, "5.3. Luong ket ban", 2)
    add_numbered(doc, [
        "Nguoi dung nhap keyword vao search box.",
        "searchUsers() goi GET /api/users/search?q=keyword.",
        "createConversationCard(type='search') render card user co ten, username va nut Ket ban.",
        "sendFriendRequestFromCard() goi POST /api/friends/requests.",
        "FriendController.SendFriendRequest() tao request Pending trong MongoDB.",
        "Server phat friend-request-created toi group user nguoi nhan; chuong thong bao cap nhat badge realtime.",
    ])

    add_image(doc, "03-signalr-events.png", "Hinh 4. Bang SignalR events cua he thong.", 6.1)

    doc.add_page_break()

    add_heading(doc, "6. API Va Response Chinh", 1)
    add_paragraph(doc, "API tra ve theo huong frontend co the render truc tiep. Rieng friend va conversation response da bo sung userName de UI dung thong nhat, van giu username de khong pha code cu.")
    add_image(doc, "04-api-map.png", "Hinh 5. API map: endpoint, controller va ham frontend goi.", 6.1)
    add_image(doc, "05-response-map.png", "Hinh 6. Cac response object quan trong.", 6.1)

    add_heading(doc, "Bang API tom tat", 2)
    add_matrix_table(
        doc,
        ["Endpoint", "Frontend goi tu ham", "Chuc nang"],
        [
            ("GET /api/users/me", "loadChatData(), Protect_Chat.js", "Lay user dang dang nhap."),
            ("GET /api/users/search", "searchUsers()", "Tim user de ket ban/xem profile."),
            ("GET /api/users/{id}/profile", "openUserProfile()", "Lay ten, SDT, bio, nickname."),
            ("GET /api/chat", "loadChatData()", "Lay danh sach conversation."),
            ("POST /api/chat", "getOrCreateDirectConversation()", "Tao direct/group conversation."),
            ("GET /api/chat/{id}/messages", "selectConversation()", "Lay lich su tin nhan."),
            ("POST /api/chat/{id}/messages", "Submit form message", "Gui tin nhan va phat realtime."),
            ("GET /api/friends", "loadChatData()", "Lay danh sach ban be."),
            ("GET /api/friends/requests", "loadFriendRequests()", "Lay loi moi ket ban."),
            ("POST /api/friends/requests", "sendFriendRequestFromCard()", "Gui loi moi ket ban."),
        ],
        [2.1, 2.2, 2.2],
    )

    add_heading(doc, "7. SignalR Realtime", 1)
    add_paragraph(
        doc,
        "SignalR duoc dung de server day su kien den client ma khong can frontend lien tuc polling API. Khi user connect, ChatHub dua connection vao group user:{id}; khi user mo conversation, client join group conversation:{id}."
    )
    add_matrix_table(
        doc,
        ["Event", "Server phat tu", "Client xu ly", "Tac dong UI"],
        [
            ("new-conversation", "ChatController.CreateConversation()", "connectRealtime()", "Them/cap nhat conversation trong sidebar."),
            ("new-message", "ChatController.SendMessage()", "connectRealtime()", "Cap nhat last message va them bubble realtime."),
            ("read-message", "ChatController.MarkAsSeen()", "connectRealtime()", "Cap nhat unread/seen."),
            ("friend-request-created", "FriendController.SendFriendRequest()", "connectRealtime()", "Tang badge chuong va them card loi moi."),
            ("friend-request-accepted", "FriendController.AcceptFriendRequest()", "connectRealtime()", "Cap nhat danh sach ban be realtime."),
            ("friend-request-declined", "FriendController.DeclineFriendRequest()", "connectRealtime()", "Xoa request khoi tab da gui."),
        ],
        [1.55, 1.85, 1.35, 1.75],
    )

    add_heading(doc, "8. Cau Truc File Chinh", 1)
    add_matrix_table(
        doc,
        ["File", "Vai tro", "Ham/phan quan trong"],
        [
            ("Frontend/Views/Chat/Index.cshtml", "Layout UI chat", "notification-bell, friendModal, friendListCard, profileModal"),
            ("Frontend/wwwroot/css/Chat.css", "Style UI chat", "theme, sidebar, card, modal, animation tai tho"),
            ("Frontend/wwwroot/js/ChatAPI/ChatAction.js", "Logic frontend", "loadChatData, searchUsers, sendFriendRequestFromCard, connectRealtime"),
            ("Frontend/wwwroot/js/ChatAPI/Api_Fetch.js", "Fetch co token", "apiFetch, refresh token khi 403"),
            ("Backend/Controllers/ChatController.cs", "API chat", "FormatConversation, CreateConversation, SendMessage"),
            ("Backend/Controllers/FriendController.cs", "API ket ban", "FormatUser, SendFriendRequest, AcceptFriendRequest, DeclineFriendRequest"),
            ("Backend/Controllers/UserController.cs", "API user", "GetMe, Search, GetProfile"),
            ("Backend/Hubs/ChatHub.cs", "Realtime hub", "OnConnectedAsync, JoinConversation, LeaveConversation"),
            ("Backend/Helper/ConversationHelper.cs", "Helper conversation", "IsParticipant, GetParticipantIds, UpdateAfterCreateMessage"),
        ],
        [2.15, 1.75, 2.6],
    )

    add_heading(doc, "9. Kiem Thu Va Ket Qua", 1)
    add_bullets(doc, [
        "Build backend/frontend ASP.NET thanh cong voi lenh dotnet build BunnyChat.csproj --no-restore.",
        "Ung dung local chay duoc tren http://localhost:5281.",
        "Trang /Chat co co che bao ve route: khi chua co token se redirect ve trang dang nhap.",
        "Graph HTML duoc mo qua /docs/system-graph.html va chup anh thanh cong.",
        "Bao cao nay duoc render/kiem tra lai bang quy trinh DOCX -> PNG truoc khi ban giao.",
    ])

    add_heading(doc, "10. Han Che Va Huong Phat Trien", 1)
    add_bullets(doc, [
        "Can bo sung tai khoan test hoac seed data rieng de demo day du chat UI ma khong phu thuoc database that.",
        "Co the them trang quan ly group: tao nhom, sua ten nhom, them/xoa thanh vien.",
        "Co the them upload anh/file cho message, typing indicator va online/offline status.",
        "Nen tach service layer cho Friend/Chat neu logic tiep tuc lon hon.",
        "Nen viet unit test va integration test cho API friend, conversation va SignalR.",
    ])

    add_heading(doc, "Ket Luan", 1)
    add_paragraph(
        doc,
        "BunnyChat hien da co nen tang chat realtime bang ASP.NET Core, MongoDB va SignalR. Phan UI chu de tho, tim kiem user, ket ban, thong bao loi moi va realtime events da duoc thiet ke de sau nay render truc tiep theo data server tra ve. He thong van giu duoc huong kien truc hien tai va co graph chi tiet de tiep tuc mo rong."
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_report()
    print(output)
