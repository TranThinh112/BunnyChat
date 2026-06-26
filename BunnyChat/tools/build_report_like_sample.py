from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from zipfile import ZipFile
import shutil
import re
import tempfile


def U(text):
    return text.encode("ascii").decode("unicode_escape")


ROOT = Path(r"C:\Users\Admin\Documents\BunnyChat")
WORKSPACE = ROOT / "BunnyChat"
SRC = ROOT / "BaoCao_DoAn_BunnyChat.docx"
OUT = WORKSPACE / "BaoCao_DoAn_BunnyChat_theo_mau.docx"
SAMPLE_COVER = ROOT / "_report_sample" / "[1] Trang bia bao cao.docx"


def set_run_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(key), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_pfmt(paragraph, align=None, before=0, after=6, line=1.5):
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_style(doc, style_name, size, bold=False, italic=False, align=None, before=0, after=6):
    style = doc.styles[style_name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(key), "Times New Roman")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.5
    if align is not None:
        style.paragraph_format.alignment = align


def replace_para_text(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=12)


def add_center(doc, text, size=13, bold=False, after=6, before=0, line=1.2):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.CENTER, after=after, line=line)
    return paragraph


def add_label_line(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(5.2)
    label_run = paragraph.add_run(label + " ")
    set_run_font(label_run, size=13, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=13)
    set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.LEFT, after=3, line=1.25)
    return paragraph


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.different_first_page_header_footer = True
    add_page_number(section.footer.paragraphs[0])

    set_style(doc, "Normal", 13, False, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6)
    set_style(doc, "Heading 1", 16, True, False, WD_ALIGN_PARAGRAPH.LEFT, 12, 8)
    set_style(doc, "Heading 2", 14, True, True, WD_ALIGN_PARAGRAPH.LEFT, 8, 6)
    set_style(doc, "Heading 3", 13, True, False, WD_ALIGN_PARAGRAPH.LEFT, 6, 4)
    try:
        set_style(doc, "List Paragraph", 13, False, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 4)
    except Exception:
        pass

    body = doc.element.body
    for idx in range(17, -1, -1):
        body.remove(body[idx])

    replacements = {
        U("lu\\u1ed3n"): U("lu\\u1ed3ng"),
        U("d\\u0103ng nh\\u1eadp"): U("\\u0111\\u0103ng nh\\u1eadp"),
        U("l\\u1eddi m\\u1edbi k\\u1ebft b\\u1ea1n"): U("l\\u1eddi m\\u1eddi k\\u1ebft b\\u1ea1n"),
        U("Chuc nang"): U("Ch\\u1ee9c n\\u0103ng"),
        U("direct user"): U("redirect user"),
        U("duoc mo"): U("\\u0111\\u01b0\\u1ee3c m\\u1edf"),
        U("ch\\u1ee7 \\u0111\\u1ec1 th\\u00f4"): U("ch\\u1ee7 \\u0111\\u1ec1 th\\u1ecf"),
        U("5.1. Lu\\u1ed3ng m\\u1edf trang chat"): U("4.2. Lu\\u1ed3ng m\\u1edf trang chat"),
        U("5.2. Lu\\u1ed3ng g\\u1eedi tin nh\\u1eafn realtime"): U("4.3. Lu\\u1ed3ng g\\u1eedi tin nh\\u1eafn realtime"),
        U("5.3. Lu\\u1ed3ng k\\u1ebft b\\u1ea1n"): U("4.4. Lu\\u1ed3ng k\\u1ebft b\\u1ea1n"),
        "Hinh ": U("H\\u00ecnh "),
        U("5.1. lu\\u1ed3ng"): U("5.1. Lu\\u1ed3ng"),
    }
    chapter_map = {
        U("1. Gi\\u1edbi thi\\u1ec7u \\u0111\\u1ed3 \\u00e1n"): U("Ch\\u01b0\\u01a1ng 1. T\\u1ed4NG QUAN"),
        U("2. C\\u00e1c ch\\u1ee9c n\\u0103ng"): U("Ch\\u01b0\\u01a1ng 3. K\\u1ebeT QU\\u1ea2 TH\\u1ef0C NGHI\\u1ec6M"),
        U("3. \\u1ea2nh minh h\\u1ecda ch\\u1ee9c n\\u0103ng"): U("3.1. \\u1ea2nh minh h\\u1ecda ch\\u1ee9c n\\u0103ng"),
        U("4. Ki\\u1ebfn tr\\u00fac t\\u1ed5ng quan"): U("Ch\\u01b0\\u01a1ng 4. KI\\u1ebeN TR\\u00daC V\\u00c0 LU\\u1ed2NG X\\u1eec L\\u00dd H\\u1ec6 TH\\u1ed0NG"),
        U("5. Lu\\u1ed3ng logic ho\\u1ea1t \\u0111\\u1ed9ng"): U("4.1. Lu\\u1ed3ng logic ho\\u1ea1t \\u0111\\u1ed9ng"),
        U("6. API, SignalR Event v\\u00e0 Response ch\\u00ednh"): U("Ch\\u01b0\\u01a1ng 5. API, SIGNALR V\\u00c0 C\\u1ea4U TR\\u00daC FILE"),
        U("6. API V\\u00e0 Response Ch\\u00ednh"): U("Ch\\u01b0\\u01a1ng 5. API, SIGNALR V\\u00c0 C\\u1ea4U TR\\u00daC FILE"),
        U("7. SignalR Realtime"): U("5.1. SignalR Realtime"),
        U("7. C\\u1ea5u tr\\u00fac file ch\\u00ednh"): U("5.1. C\\u1ea5u tr\\u00fac file ch\\u00ednh"),
        U("8. C\\u1ea5u tr\\u00fac c\\u00e1c file ch\\u00ednh"): U("5.2. C\\u1ea5u tr\\u00fac c\\u00e1c file ch\\u00ednh"),
        U("8. K\\u1ebft lu\\u1eadn v\\u00e0 h\\u01b0\\u1edbng ph\\u00e1t tri\\u1ec3n"): U("Ch\\u01b0\\u01a1ng 6. K\\u1ebeT LU\\u1eacN V\\u00c0 H\\u01af\\u1edaNG PH\\u00c1T TRI\\u1ec2N"),
        U("9. Ki\\u1ec3m th\\u1eed v\\u00e0 k\\u1ebft qu\\u1ea3"): U("5.3. Ki\\u1ec3m th\\u1eed v\\u00e0 k\\u1ebft qu\\u1ea3"),
        U("K\\u1ebft lu\\u1eadn"): U("Ch\\u01b0\\u01a1ng 6. K\\u1ebeT LU\\u1eacN V\\u00c0 H\\u01af\\u1edaNG PH\\u00c1T TRI\\u1ec2N"),
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        stripped = new_text.strip()
        if stripped in chapter_map:
            new_text = chapter_map[stripped]
            stripped = new_text
        if new_text != text:
            replace_para_text(paragraph, new_text)

        if stripped.startswith(U("Ch\\u01b0\\u01a1ng ")):
            paragraph.style = "Heading 1"
        elif re.match(r"^\d+\.\d+\.", stripped):
            paragraph.style = "Heading 2"
        elif stripped.startswith(U("H\\u00ecnh ")) or stripped.startswith(U("B\\u1ea3ng ")):
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif paragraph.style.name.startswith("Heading"):
            paragraph.style = "Normal"

        for run in paragraph.runs:
            if stripped.startswith(U("H\\u00ecnh ")) or stripped.startswith(U("B\\u1ea3ng ")):
                set_run_font(run, size=12, italic=True)
            elif paragraph.style.name == "Heading 1":
                set_run_font(run, size=16, bold=True, italic=False)
            elif paragraph.style.name == "Heading 2":
                set_run_font(run, size=14, bold=True, italic=True)
            else:
                set_run_font(run, size=13)

        if paragraph.style.name == "Heading 1":
            set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.LEFT, before=12, after=8)
        elif paragraph.style.name == "Heading 2":
            set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)
        elif stripped.startswith(U("H\\u00ecnh ")) or stripped.startswith(U("B\\u1ea3ng ")):
            set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.CENTER, after=8)
        else:
            set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY, after=6)

    body = doc.element.body
    for i, element in enumerate(list(body)):
        text = "".join(t.text for t in element.iter(qn("w:t")) if t.text).strip()
        if text == U("Ch\\u01b0\\u01a1ng 3. K\\u1ebeT QU\\u1ea2 TH\\u1ef0C NGHI\\u1ec6M"):
            heading = doc.add_paragraph(U("Ch\\u01b0\\u01a1ng 2. C\\u01a0 S\\u1ede L\\u00dd THUY\\u1ebeT V\\u00c0 C\\u00d4NG NGH\\u1ec6 S\\u1eec D\\u1ee4NG"))
            heading.style = "Heading 1"
            paragraph = doc.add_paragraph(
                U("\\u0110\\u1ed3 \\u00e1n s\\u1eed d\\u1ee5ng ASP.NET Core MVC \\u0111\\u1ec3 render giao di\\u1ec7n, Web API \\u0111\\u1ec3 x\\u1eed l\\u00fd nghi\\u1ec7p v\\u1ee5, SignalR \\u0111\\u1ec3 truy\\u1ec1n d\\u1eef li\\u1ec7u th\\u1eddi gian th\\u1ef1c v\\u00e0 MongoDB \\u0111\\u1ec3 l\\u01b0u tr\\u1eef ng\\u01b0\\u1eddi d\\u00f9ng, b\\u1ea1n b\\u00e8, h\\u1ed9i tho\\u1ea1i, tin nh\\u1eafn. Frontend d\\u00f9ng HTML, CSS v\\u00e0 JavaScript thu\\u1ea7n \\u0111\\u1ec3 g\\u1ecdi API, c\\u1eadp nh\\u1eadt giao di\\u1ec7n v\\u00e0 nh\\u1eadn s\\u1ef1 ki\\u1ec7n realtime t\\u1eeb server.")
            )
            paragraph.style = "Normal"
            for para in (heading, paragraph):
                for run in para.runs:
                    set_run_font(run, size=16 if para.style.name == "Heading 1" else 13, bold=True if para.style.name == "Heading 1" else None)
            heading_el = heading._p
            paragraph_el = paragraph._p
            body.remove(heading_el)
            body.remove(paragraph_el)
            body.insert(i, paragraph_el)
            body.insert(i, heading_el)
            break

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        set_run_font(run, size=12, bold=True if row_idx == 0 else None)
                if row_idx == 0:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = tc_pr.find(qn("w:shd"))
                    if shd is None:
                        shd = OxmlElement("w:shd")
                        tc_pr.append(shd)
                    shd.set(qn("w:fill"), "EDE7F6")

    max_width = Cm(15.8)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = int(shape.width * ratio)
            shape.height = int(shape.height * ratio)

    logo_path = None
    try:
        tmpdir = Path(tempfile.mkdtemp(prefix="bunny_report_"))
        with ZipFile(SAMPLE_COVER) as archive:
            media = [name for name in archive.namelist() if name.startswith("word/media/")]
            if media:
                logo_path = tmpdir / Path(media[0]).name
                logo_path.write_bytes(archive.read(media[0]))
    except Exception:
        logo_path = None

    start_marker = len(doc.element.body) - 1
    add_center(doc, U("B\\u1ed8 GI\\u00c1O D\\u1ee4C V\\u00c0 \\u0110\\u00c0O T\\u1ea0O"), 13, True, 0)
    add_center(doc, U("TR\\u01af\\u1edcNG \\u0110\\u1ea0I H\\u1eccC C\\u00d4NG NGH\\u1ec6 TP. HCM"), 13, True, 0)
    add_center(doc, U("KHOA C\\u00d4NG NGH\\u1ec6 TH\\u00d4NG TIN"), 13, True, 18)
    if logo_path and logo_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo_path), width=Cm(3.0))
        set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.CENTER, after=20, line=1.0)
    else:
        add_center(doc, "", 13, False, 24)
    add_center(doc, U("B\\u00c1O C\\u00c1O \\u0110\\u1ed2 \\u00c1N"), 20, True, 10)
    add_center(doc, U("BUNNYCHAT - \\u1ee8NG D\\u1ee4NG CHAT REALTIME"), 18, True, 8)
    add_center(doc, "ASP.NET CORE MVC + WEB API + SIGNALR + MONGODB", 14, True, 26)
    add_label_line(doc, U("Ng\\u00e0nh:"), U("C\\u00f4ng ngh\\u1ec7 th\\u00f4ng tin"))
    add_label_line(doc, U("T\\u00ean h\\u1ecdc ph\\u1ea7n:"), U("L\\u1eadp tr\\u00ecnh Web"))
    add_label_line(doc, U("Gi\\u1ea3ng vi\\u00ean h\\u01b0\\u1edbng d\\u1eabn:"), U("Tr\\u1ecbnh \\u0110\\u1ed3ng Th\\u1ea1ch Tr\\u00fac"))
    add_label_line(doc, U("Sinh vi\\u00ean th\\u1ef1c hi\\u1ec7n:"), U("Tr\\u1ea7n T\\u0103ng Tr\\u01b0\\u1eddng Th\\u1ecbnh"))
    add_label_line(doc, "MSSV:", "2387701939")
    add_center(doc, U("TP. H\\u1ed3 Ch\\u00ed Minh, 2026"), 13, True, 0, 38)
    doc.add_page_break()
    add_center(doc, U("M\\u1ee4C L\\u1ee4C T\\u00d3M T\\u1eaeT"), 16, True, 8)
    for item in [
        U("Ch\\u01b0\\u01a1ng 1. T\\u1ed4NG QUAN"),
        U("Ch\\u01b0\\u01a1ng 2. C\\u01a0 S\\u1ede L\\u00dd THUY\\u1ebeT V\\u00c0 C\\u00d4NG NGH\\u1ec6 S\\u1eec D\\u1ee4NG"),
        U("Ch\\u01b0\\u01a1ng 3. K\\u1ebeT QU\\u1ea2 TH\\u1ef0C NGHI\\u1ec6M"),
        U("Ch\\u01b0\\u01a1ng 4. KI\\u1ebeN TR\\u00daC V\\u00c0 LU\\u1ed2NG X\\u1eec L\\u00dd H\\u1ec6 TH\\u1ed0NG"),
        U("Ch\\u01b0\\u01a1ng 5. API, SIGNALR V\\u00c0 C\\u1ea4U TR\\u00daC FILE"),
        U("Ch\\u01b0\\u01a1ng 6. K\\u1ebeT LU\\u1eacN V\\u00c0 H\\u01af\\u1edaNG PH\\u00c1T TRI\\u1ec2N"),
    ]:
        paragraph = doc.add_paragraph(item)
        paragraph.style = "Normal"
        paragraph.paragraph_format.left_indent = Cm(1)
        for run in paragraph.runs:
            set_run_font(run, size=13)
        set_pfmt(paragraph, WD_ALIGN_PARAGRAPH.LEFT, after=4)
    doc.add_page_break()

    body = doc.element.body
    all_elements = list(body)
    front_elements = [element for element in all_elements[start_marker:] if element.tag != qn("w:sectPr")]
    for element in front_elements:
        body.remove(element)
    for index, element in enumerate(front_elements):
        body.insert(index, element)

    doc.core_properties.title = U("B\\u00e1o c\\u00e1o \\u0111\\u1ed3 \\u00e1n BunnyChat")
    doc.core_properties.subject = U("\\u1ee8ng d\\u1ee5ng chat realtime ASP.NET Core SignalR MongoDB")
    doc.core_properties.author = U("Tr\\u1ea7n T\\u0103ng Tr\\u01b0\\u1eddng Th\\u1ecbnh")
    doc.core_properties.keywords = "BunnyChat, ASP.NET Core, SignalR, MongoDB"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
